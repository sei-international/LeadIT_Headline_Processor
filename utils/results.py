"""
This module provides functions to generate and format output documents for the GPT Batch Policy Processor.
It includes functions to create tables in Word documents, read data from Excel files, and format the output
documents with relevant information and metrics.

Functions:
- get_output_fname: Generates the output file name based on the provided path function and file type.
- create_word_table: Creates a table in a Word document with the provided data.
- format_output_doc: Formats the output Word document with query and variable specifications.
- output_results: Outputs the results to a Word document.
- output_metrics: Outputs processing metrics to a Word document.
"""

from datetime import datetime
from docx.shared import Pt
import os
import io
import pandas as pd
from utils.validate_results import get_check_results_flag
import streamlit as st
from services.onedrive import get_graph_api_token, upload_file_to_onedrive
# Function to generate the output file name based on the provided path function and file type
def get_output_fname(path_fxn, filetype="xlsx"):
    return path_fxn(f"results.{filetype}")

def create_word_doc(doc, output_pdf_path, rows_dict):
    fname = os.path.basename(output_pdf_path)
    doc.add_heading(f"{fname}", 2)
    for title in rows_dict[1]: 
        doc.add_paragraph(title)


def output_results(gpt_analyzer, output_doc, output_pdf_path, headline_info):
    # Extract relevant data from headline_info before passing it
    processed_info = []
    for entry in headline_info:
        title = entry.get('title', 'N/A')
        relevance = entry.get('relevant', 'no')
        # Convert relevance dictionary into a readable string
        processed_info.append({
            "Title": title,
            "Relevant": relevance
        })

    rows_dict = gpt_analyzer.get_results(processed_info)
    print("Rows Dict:", rows_dict)
    create_word_doc(output_doc, output_pdf_path, rows_dict)


def output_metrics(doc, num_docs, t, num_pages, failed_pdfs):
    doc.add_heading(
        f"{num_docs} documents ({num_pages} total pages) processed in {t:.2f} seconds",
        4,
    )
    if len(failed_pdfs) > 0:
        doc.add_heading(f"Unable to process the following PDFs: {failed_pdfs}", 4)
def output_results_excel(relevant_articles, irrelevant_articles, output_path):
    """
    Writes the results into an Excel file with four worksheets:
      - 'Relevant Stage 1': Articles flagged as relevant by the headline but with insufficient extracted core details.
           (Contains only the article title and URL.)
      - 'Relevant Stage 2': Articles flagged as relevant by the headline that contain at least a project name and a company.
           (Includes extra columns such as Project name, Project scale, Year to be online, Technology to be used,
            Company, Potential Partners, Continent, Country, Project status, and a "Check Results" column.)
      - 'Irrelevant': Articles deemed irrelevant.
      - 'All Articles': A combined list of all articles (from irrelevant, Stage 1, and Stage 2) showing
           the article titles, URLs, and if they were discarded before stage 1 or stage 2.
    """

    # Define simple columns for Stage 1 and Irrelevant sheets.
    simple_cols = ["title", "url"]

    # Define detailed columns for the Relevant Stage 2 sheet.
    detailed_cols = [
        "Internal ID", 
        "Justification",
        "Project name",
        "Project scale",
        "Year to be online",
        "Technology to be used",
        "Company",
        "Potential Partners",
        "Company type",
        "Project type",
        "Company has climate goals?",
        "Production plant",
        "Updated GEM Plant ID",
        "GEM wiki page link",
        "Latitude",
        "Longitude",
        "Coordinate accuracy",
        "Continent",
        "Country",
        "Iron production capacity (million tonnes per year)",
        "Steel production capacity (million tonnes per year)",
        "States iron & steel capacity?",
        "[ref] Iron or steel capacity",
        "Hydrogen generation capacity (MW)",
        "States CC & H2 capacity?",
        "[ref] CC or H2 capacity",
        "[ref] Investment",
        "Business proposed",
        "Project status",
        "Year construction began",
        "Actual start year",
        "[ref] Date of announcement",
        "Comments",
        "Lastest project news (yyyy-mm-dd)",
        "Lastly updated (yyyy-mm-dd)",
        "References 1",
        "Reference Article",
        "Check Results"  # New column for validation flag
    ]

    # First, filter out articles that have been flagged as irrelevant via the "irrelevant" key.
    newly_irrelevant = []
    filtered_relevant_articles = []
    for article in relevant_articles:
        if article.get("irrelevant"):  # If flagged as irrelevant by the secondary query.
            newly_irrelevant.append(article)
        else:
            filtered_relevant_articles.append(article)
    print(newly_irrelevant)
    # Append newly flagged articles to the existing irrelevant_articles list.
    irrelevant_articles.extend(newly_irrelevant)

    # Now, separate the remaining relevant articles into Stage 1 and Stage 2.
    # Stage 2 requires that both 'company' and 'project_name' are non-empty.
    stage1_articles = []  # Articles with insufficient details (i.e. did not reach Stage 2).
    stage2_articles = []  # Articles with both 'company' and 'project_name' provided.
    for article in filtered_relevant_articles:
        if article.get("company", "").strip() and article.get("project_name", "").strip():
            stage2_articles.append(article)
        else:
            stage1_articles.append(article)

    print("Stage 1 articles:", stage1_articles)
    print("Stage 2 articles:", stage2_articles)

    # Build DataFrame for "Relevant Stage 1" (only title and URL).
    df_stage1 = pd.DataFrame(stage1_articles, columns=simple_cols)

    # Helper function to build a detailed row for a Stage 2 article.
    def build_detailed_row(article):
        row_data = {col: "" for col in detailed_cols}
        # Core details.
        row_data["Project name"] = article.get("project_name", "")
        row_data["Project scale"] = article.get("scale", "")
        row_data["Year to be online"] = article.get("timeline", "")
        row_data["Technology to be used"] = article.get("technology", "")
        # Additional details.
        row_data["Company"] = article.get("company", "")
        row_data["Potential Partners"] = article.get("partners", "")
        row_data["Continent"] = article.get("continent", "")
        row_data["Country"] = article.get("country", "")
        row_data["Project status"] = article.get("project_status", "")
        # Basic article info.
        row_data["Reference Article"] = article.get("title", "")
        row_data["References 1"] = article.get("url", "")
        
        # Use the full article text (if available) for fuzzy-checking.
        print("full text", article.get("full_text"))
        if "full_text" in article and article["full_text"]:
            core_extracted = {
                "project_name": article.get("project_name", ""),
                "scale": article.get("scale", ""),
                "timeline": article.get("timeline", ""),
                "technology": article.get("technology", "")
            }
            flag, scores = get_check_results_flag(core_extracted, article["full_text"])
            row_data["Check Results"] = flag
        else:
            row_data["Check Results"] = ""
        return row_data

    # Build DataFrame for "Relevant Stage 2".
    detailed_rows = [build_detailed_row(article) for article in stage2_articles]
    print("Detailed rows:", detailed_rows)
    df_stage2 = pd.DataFrame(detailed_rows, columns=detailed_cols)

    # Build DataFrame for "Irrelevant" articles.
    df_irrelevant = pd.DataFrame(irrelevant_articles, columns=simple_cols)
    print("Irrelevant DataFrame:", df_irrelevant)
    print("Detailed DataFrame:", df_stage2)
    print("Simple DataFrame:", df_stage1)

    # Build DataFrame for "All Articles"
    # It will contain title, URL, and a "Discarded" column indicating the discard stage.
    all_articles = []
    # For Stage 1 articles, mark as "Discarded before Stage 2".
    for article in stage1_articles:
        all_articles.append({
            "title": article.get("title", ""),
            "url": article.get("url", ""),
            "Discarded": "Discarded before Stage 2"
        })
    # For Stage 2 articles, leave the "Discarded" column blank.
    for article in stage2_articles:
        all_articles.append({
            "title": article.get("title", ""),
            "url": article.get("url", ""),
            "Discarded": ""
        })
    # For irrelevant articles, mark as "Discarded before Stage 1".
    for article in irrelevant_articles:
        all_articles.append({
            "title": article.get("title", ""),
            "url": article.get("url", ""),
            "Discarded": "Discarded before Stage 1"
        })
    df_all = pd.DataFrame(all_articles, columns=["title", "url", "Discarded"])
    tenant_id = st.secrets["od_tenantid"]
    client_id = st.secrets["od_client_id"]
    client_secret = st.secrets["od_client_value"]
    drive_id = st.secrets["od_drive_id"]
    parent_item_id = st.secrets["od_parent_item"]
    buffer = io.BytesIO()
    # Write all DataFrames to an Excel file with four sheets.
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_stage1.to_excel(writer, sheet_name="Relevant Stage 1", index=False)
        df_stage2.to_excel(writer, sheet_name="Relevant Stage 2", index=False)
        df_irrelevant.to_excel(writer, sheet_name="Irrelevant", index=False)
        df_all.to_excel(writer, sheet_name="All Articles", index=False)
    buffer.seek(0)
    file_bytes = buffer.getvalue()
    graph_access_token = get_graph_api_token(tenant_id, client_id, client_secret)
    if not graph_access_token:
        print("Could not obtain Graph API token.")
        return
    upload_file_to_onedrive(file_bytes, drive_id, parent_item_id, output_path, graph_access_token)
    print(f"Excel saved to: {output_path}")
